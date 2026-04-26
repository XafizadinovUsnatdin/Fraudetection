"""SafeNet — Model tushuntiruvchi Word hujjati (.docx) generatori."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = Path(__file__).parent / "SafeNet_Model_Tushuntirish.docx"

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(doc, text, bold=False, size=11, color_hex=None):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    return para

def bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run  = para.add_run(text)
    run.font.size = Pt(11)
    return para

def numbered(doc, text, level=0):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run  = para.add_run(text)
    run.font.size = Pt(11)
    return para

def divider(doc):
    doc.add_paragraph("─" * 80)

def simple_table(doc, headers, rows,
                 header_bg="1E3A5F", header_fg="FFFFFF",
                 alt_bg="EBF2FF", col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_bg(cell, header_bg)
        cell_text(cell, h, bold=True, size=10,
                  color=header_fg, align=WD_ALIGN_PARAGRAPH.CENTER)
    # rows
    for i, row_data in enumerate(rows):
        bg = alt_bg if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            set_cell_bg(cell, bg)
            cell_text(cell, str(val), size=10)
    # column widths
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ══════════════════════════════════════════════════════════════════════════
    # SARLAVHA
    # ══════════════════════════════════════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("SafeNet — Anti-Fraud AI Tizimi")
    run.bold = True; run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run("DeepFraudNet Modeli: To'liq Texnik Tushuntirish")
    run2.font.size = Pt(14); run2.italic = True

    sub_para2 = doc.add_paragraph()
    sub_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = sub_para2.add_run("Mualliflar: Xafizadinov Usnatdin  |  2026 yil")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. KIRISH
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "1. Kirish — Muammo va Maqsad")
    body(doc,
         "Moliyaviy firibgarlik (fraud) butun dunyo bo'ylab yiliga 485 milliard dollar zarar "
         "keltiradi. O'zbekistonda raqamli to'lovlar o'sishi bilan birga shubhali tranzaksiya "
         "holatlari ham ko'paymoqda. SafeNet loyihasi — chuqur o'qitish asosida qurilgan "
         "real-vaqt anti-fraud tizimi bo'lib, har bir to'lov so'rovini 100 millisekunddan "
         "kam vaqt ichida baholaydi va uch qarordan birini qaytaradi: ALLOW, REVIEW yoki BLOCK.")

    doc.add_paragraph()
    body(doc, "Tizimning asosiy imkoniyatlari:", bold=True)
    bullet(doc, "25 ta belgi asosida chuqur neyron tarmoq (DeepFraudNet) orqali baholash")
    bullet(doc, "Tranzaksiya vaqti, qurilma, joylashuv va miqdor bo'yicha risk signallari")
    bullet(doc, "Interaktiv Streamlit demo paneli — jonli simulyatsiya rejimi")
    bullet(doc, "REST API orqali bank tizimlariga ulanish imkoniyati")
    bullet(doc, "50,000 sintetik tranzaksiyada o'qitilgan, ROC-AUC 0.857")

    # ══════════════════════════════════════════════════════════════════════════
    # 2. DATASET
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "2. Dataset Ko'rinishi va Tuzilishi")
    body(doc,
         "Model 50,000 qatordan iborat sintetik tranzaksiyalar to'plamida o'qitilgan. "
         "Dataset Kaggle'dan olingan real namunalar asosida generatsiya qilingan bo'lib, "
         "5% fraud ulushiga ega (klasslar nomutanosibligi — sinf muvozanati muammosi).")

    doc.add_paragraph()
    body(doc, "Dataset asosiy ko'rsatkichlari:", bold=True)
    simple_table(doc,
        ["Ko'rsatkich", "Qiymat"],
        [
            ["Jami qatorlar", "50,000"],
            ["Fraud qatorlar (%)", "2,500 qator (5.0%)"],
            ["Normal qatorlar (%)", "47,500 qator (95.0%)"],
            ["Ustunlar soni", "23 ta (model uchun 18 ta ishlatiladi)"],
            ["Foydalanuvchilar", "~1,000 noyob user_id"],
            ["Joylashuvlar", "Toshkent, Samarqand, Buxoro, Farg'ona, Namangan, foreign_ip, unknown va b."],
            ["Qurilmalar", "ios, android, chrome, safari, emulator, linux, unknown"],
            ["Tranzaksiya soatlari", "0–23 (tun soati 0–5 xavfli hisoblanadi)"],
        ],
        col_widths=[2.5, 4.0]
    )

    doc.add_paragraph()
    body(doc, "Asosiy ustunlar va ma'nosi:", bold=True)
    simple_table(doc,
        ["Ustun nomi", "Turi", "Ma'nosi"],
        [
            ["amount", "Raqamli", "Tranzaksiya summasi (so'm)"],
            ["transaction_hour", "Raqamli (0–23)", "Tranzaksiya soati"],
            ["Transaction Frequency", "Raqamli", "Foydalanuvchi tranzaksiya chastotasi"],
            ["Time Since Last Transaction", "Raqamli", "Oxirgi tranzaksiyadan o'tgan vaqt (soat)"],
            ["Account Age", "Raqamli", "Hisob yoshi (kun)"],
            ["Normalized Transaction Amount", "Raqamli (0–1)", "Normalangan summa"],
            ["Fraud Complaints Count", "Raqamli", "Foydalanuvchiga bog'liq shikoyatlar soni"],
            ["Geo-Location Flags", "Ikkilik (0/1)", "Joylashuv anomaliyasi bayrog'i"],
            ["Location-Inconsistent Transactions", "Ikkilik", "Joylashuv nomuvofiqlik bayrog'i"],
            ["Recipient Verification Status", "Ikkilik", "Qabul qiluvchi tasdiqlanganmi"],
            ["Recipient Blacklist Status", "Ikkilik", "Qabul qiluvchi qora ro'yxatdami"],
            ["VPN or Proxy Usage", "Ikkilik", "VPN/Proxy ishlatilganmi"],
            ["Merchant Category Mismatch", "Ikkilik", "Savdo kategoriyasi mos kelmayapti"],
            ["User Daily Limit Exceeded", "Ikkilik", "Kunlik limit oshib ketdimi"],
            ["Recent High-Value Transaction Flags", "Ikkilik", "Yaqinda yuqori summa bayrog'i"],
            ["Past Fraudulent Behavior Flags", "Ikkilik", "Oldingi fraud xatti-harakatlar"],
            ["device", "Kategorik", "Qurilma turi: ios, android, emulator va h.k."],
            ["location", "Kategorik", "Joylashuv: Toshkent, foreign_ip va h.k."],
            ["isFraud", "Maqsad (0/1)", "Fraud=1, Normal=0"],
        ],
        col_widths=[2.8, 1.6, 3.0]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 3. MA'LUMOTLARNI TAYYORLASH
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "3. Ma'lumotlarni Tayyorlash (Preprocessing)")
    body(doc,
         "Xom dataset modelga berilishdan oldin quyidagi ketma-ket qadamlardan o'tadi:")

    doc.add_paragraph()
    numbered(doc, "Ma'lumot yuklash — CSV fayl o'qiladi, majburiy ustunlar tekshiriladi.")
    numbered(doc, "Tozalash — NaN qiymatlar 0 bilan to'ldiriladi; manfiy summalar klip qilinadi; "
                  "transaction_hour 0–23 oralig'iga cheklanadi; isFraud {0,1} ga qaytariladi.")
    numbered(doc, "Stratified namuna — 20,000 qator stratified sampling bilan tanlanadi "
                  "(fraud ulushi saqlanib qoladi).")
    numbered(doc, "Interaction features yaratish — add_interactions() funksiyasi chaqiriladi "
                  "(batafsil §4 bo'limda).")
    numbered(doc, "Train/Test ajratish — 80% o'qitish, 20% test (stratify=isFraud, seed=42).")
    numbered(doc, "Fit schema — o'qitish to'plamida raqamli ustunlar uchun mean/std hisoblanadi; "
                  "kategorik ustunlar uchun noyob qiymatlar ro'yxati saqlanadi; "
                  "NUMERICAL ustunlar uchun quantile bin chegaralari aniqlanadi.")
    numbered(doc, "Transform — test va yangi ma'lumotlar bir xil sxema bo'yicha o'zgartiriladi "
                  "(train/test leakage yo'q).")

    doc.add_paragraph()
    body(doc, "Raqamli ustunlarni normalizatsiya qilish formulasi:", bold=True)
    body(doc, "   X_norm = (X − mean) / std")
    body(doc, "Bu Z-score normalizatsiyasi. Model raqamli ustunlarni bir xil masshtabga keltirish "
             "uchun ishlatadi. Agar std = 0 bo'lsa, 1.0 bilan almashtiriladi.")

    doc.add_paragraph()
    body(doc, "Kategorik ustunlarni kodlash (OneHot Encoding):", bold=True)
    body(doc, "   device → [device_ios, device_android, device_chrome, ...]\n"
             "   location → [location_toshkent, location_samarkand, ...]")
    body(doc, "Noma'lum qiymatlar (test da uchramagan) 0 vektori bilan almashtiriladi. "
             "Bu overfitting va 'KeyError' larning oldini oladi.")

    doc.add_paragraph()
    body(doc, "Binning (miqdoriy bo'linish):", bold=True)
    body(doc, "NUMERICAL har ustun uchun 9 ta kvantil chegarasi hisoblanib, har qiymat uchun "
             "qaysi bin ichida ekanligi aniqlanadi va one-hot kodlanadi. Bu modelga "
             "nochiziqli munosabatlarni yaxshiroq o'rganishga yordam beradi.")

    # ══════════════════════════════════════════════════════════════════════════
    # 4. FEATURE ENGINEERING
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "4. Feature Engineering — Belgilar Yaratish")
    body(doc,
         "Model 18 ta asosiy belgi ustiga 7 ta qo'shimcha 'interaction' belgi qo'shilgan "
         "25 ta kengaytirilgan belgi to'plamida (EXT_FEATURES) ishlaydi.")

    doc.add_paragraph()
    body(doc, "7 ta Interaction Feature (add_interactions funksiyasi):", bold=True)
    simple_table(doc,
        ["Belgi nomi", "Formulasi", "Ma'nosi"],
        [
            ["risky_device",    "device ∈ {emulator, unknown, linux} → 1 else 0",
             "Xavfli qurilma belgisi"],
            ["risky_location",  "location ∈ {foreign_ip, unknown} → 1 else 0",
             "Xavfli joylashuv belgisi"],
            ["night_flag",      "transaction_hour ∈ [0,5] → 1 else 0",
             "Tun soati belgisi (00:00–05:59)"],
            ["dev_x_night",     "risky_device × night_flag",
             "Xavfli qurilma + tun birgaligi"],
            ["dev_x_loc",       "risky_device × risky_location",
             "Xavfli qurilma + xavfli lokatsiya"],
            ["loc_x_night",     "risky_location × night_flag",
             "Xavfli lokatsiya + tun birgaligi"],
            ["vpn_x_blacklist", "VPN or Proxy Usage × Recipient Blacklist Status",
             "VPN + qora ro'yxat birgaligi"],
        ],
        col_widths=[1.8, 2.8, 2.8]
    )

    doc.add_paragraph()
    body(doc,
         "Nima uchun interaction features muhim?",
         bold=True)
    body(doc,
         "Oddiy neyron tarmoq alohida belgilardagi nochiziqliklarni o'rganishi mumkin. "
         "Lekin 'emulator + tun soati' kombinatsiyasi alohida alohida berilgandagiga "
         "qaraganda fraud uchun ancha kuchliroq signal. Interaction features modelga "
         "ushbu kombinatsiyalarni bevosita beradi — o'rganish osonlashadi va ROC-AUC "
         "taxminan 1–2 % oshadi.")

    # ══════════════════════════════════════════════════════════════════════════
    # 5. MODEL ARXITEKTURASI
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "5. DeepFraudNet — Model Arxitekturasi")
    body(doc,
         "DeepFraudNet — bu sklearn ishlatmaydigan, to'liq numpy/pandas asosida "
         "yozilgan 3 ta yashirin qatlamli ko'p qatlamli perseptron (MLP). "
         "Streamlit Cloud'da sklearn yuklanmaydi, shuning uchun barcha matematik "
         "amallar (matritsa ko'paytmasi, gradientlar, Adam yangilanishi) qo'lda "
         "implementatsiya qilingan.")

    doc.add_paragraph()
    body(doc, "Tarmoq sxemasi:", bold=True)
    body(doc,
         "   Input (n_features)  →  Dense(256, ReLU)  →  Dense(128, ReLU)\n"
         "                        →  Dense(64, ReLU)   →  Dense(1, Sigmoid)\n"
         "\n"
         "   n_features ≈ 120–150  (raqamli + binning + OHE + interaction)")

    doc.add_paragraph()
    body(doc, "Har qatlamning vazifasi:", bold=True)
    simple_table(doc,
        ["Qatlam", "Neyronlar", "Aktivatsiya", "Vazifasi"],
        [
            ["Kirish", "~130", "—",       "Xom feature vektori"],
            ["Yashirin 1", "256", "ReLU", "Asosiy nochiziqli munosabatlarni o'rganadi"],
            ["Yashirin 2", "128", "ReLU", "Kombinatsiyalarni yanada chuqurlashtiradi"],
            ["Yashirin 3", "64",  "ReLU", "Muhim pattern'larni yanada siqadi"],
            ["Chiqish",    "1",   "Sigmoid", "P(fraud) ∈ [0, 1] — ehtimollik"],
        ],
        col_widths=[1.4, 1.2, 1.3, 3.5]
    )

    doc.add_paragraph()
    body(doc, "He initialization (og'irlik initsializatsiyasi):", bold=True)
    body(doc, "   W ~ N(0, sqrt(2 / fan_in))")
    body(doc,
         "ReLU aktivatsiyasi bilan ishlatiladigan eng yaxshi boshlang'ich og'irlik metodi. "
         "Bu gradient yo'qolishi (vanishing gradient) muammosining oldini oladi.")

    doc.add_paragraph()
    body(doc, "Forward pass (oldinga o'tish) formulasi:", bold=True)
    body(doc,
         "   Z¹ = X · W¹ + b¹\n"
         "   A¹ = ReLU(Z¹) = max(0, Z¹)\n"
         "   Z² = A¹ · W² + b²\n"
         "   A² = ReLU(Z²)\n"
         "   Z³ = A² · W³ + b³\n"
         "   A³ = ReLU(Z³)\n"
         "   Z⁴ = A³ · W⁴ + b⁴\n"
         "   ŷ  = Sigmoid(Z⁴) = 1 / (1 + e^(−Z⁴))\n"
         "   P(fraud) = ŷ")

    # ══════════════════════════════════════════════════════════════════════════
    # 6. O'QITISH JARAYONI
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "6. O'qitish Jarayoni")

    heading(doc, "6.1. Sinf Og'irligi (Class Weighting)", level=2)
    body(doc,
         "Dataset 95% normal va 5% fraud. Og'irliksiz o'qitilsa model hamma narsani "
         "normal deydi (95% aniqlik, lekin 0% fraud aniqlaydi). Bu muammoni hal qilish uchun "
         "fraud sinfiga katta og'irlik beriladi.")
    doc.add_paragraph()
    body(doc, "Og'irlik formulasi:", bold=True)
    body(doc,
         "   ratio = min( sqrt(n_negatif / n_pozitif),  5.0 )\n"
         "   sample_weight[i] = ratio  (fraud uchun)\n"
         "   sample_weight[i] = 1.0   (normal uchun)\n"
         "\n"
         "   Misol: 20,000 namuna, 1,000 fraud:\n"
         "   ratio = min( sqrt(19000 / 1000), 5.0 ) = min(4.36, 5.0) = 4.36\n"
         "\n"
         "   Nima uchun sqrt? Oddiy n/(2*n_pos) ≈ 9.5x juda kuchli bo'lib,\n"
         "   model faqat fraudni bashorat qiladi (precision <20%). sqrt 4.36x\n"
         "   muvozanatli — ham recall yuqori, ham precision qabul qilinadi.")

    doc.add_paragraph()
    heading(doc, "6.2. Yo'qotish funksiyasi (Loss Function)", level=2)
    body(doc, "   Binary Cross-Entropy (BCE) + sinf og'irligi:")
    body(doc,
         "   L = − (1/n) × Σᵢ  wᵢ × [ yᵢ × log(ŷᵢ + ε) + (1−yᵢ) × log(1−ŷᵢ + ε) ]\n"
         "\n"
         "   Bu yerda:\n"
         "     yᵢ   — haqiqiy qiymat (0 yoki 1)\n"
         "     ŷᵢ   — model bashorati (0..1)\n"
         "     wᵢ   — sinf og'irligi (fraud=4.36, normal=1.0)\n"
         "     ε    = 1e-9  (log(0) dan himoya)")

    doc.add_paragraph()
    heading(doc, "6.3. Orqaga tarqalish (Backpropagation)", level=2)
    body(doc,
         "Gradientlar zanjir qoidasi (chain rule) bilan hisoblanadi va har qatlam "
         "og'irliklari yangilanadi:")
    body(doc,
         "   Chiqish qatlami delta:\n"
         "     δ⁴ = (ŷ − y) × w  / n           ← sigmoid + BCE birgaligi soddalashadi\n"
         "\n"
         "   Yashirin qatlam delta (ReLU gradient):\n"
         "     δˡ = (δˡ⁺¹ · Wˡ⁺¹ᵀ) × (Zˡ > 0)   ← ReLU grad = 1 agar Z>0, aks 0\n"
         "\n"
         "   Og'irlik gradienti:\n"
         "     ∂L/∂Wˡ = Aˡ⁻¹ᵀ · δˡ + λ × Wˡ     ← L2 regularizatsiya qo'shildi\n"
         "     ∂L/∂bˡ = Σ δˡ")

    doc.add_paragraph()
    heading(doc, "6.4. Adam Optimizer", level=2)
    body(doc,
         "Oddiy gradient descent o'rniga Adam (Adaptive Moment Estimation) ishlatiladi. "
         "Adam har og'irlik uchun alohida o'rganish tezligini moslaydi:")
    body(doc,
         "   β₁ = 0.9   (1-momenti uchun eksponensial o'rtacha)\n"
         "   β₂ = 0.999 (2-momenti uchun eksponensial o'rtacha)\n"
         "   lr = 0.001 (o'rganish tezligi)\n"
         "   ε  = 1e-8  (bo'linishdan himoya)\n"
         "\n"
         "   m[t] = β₁ × m[t-1] + (1 − β₁) × g[t]       ← 1-moment (gradient o'rtachasi)\n"
         "   v[t] = β₂ × v[t-1] + (1 − β₂) × g[t]²      ← 2-moment (dispersiya)\n"
         "   m̂ = m[t] / (1 − β₁ᵗ)                        ← bias korreksiya\n"
         "   v̂ = v[t] / (1 − β₂ᵗ)                        ← bias korreksiya\n"
         "   W[t] = W[t-1] − lr × m̂ / (sqrt(v̂) + ε)     ← yangilanish\n"
         "\n"
         "Adam afzalligi: barqaror konvergentsiya, gradient oshib ketishi kamroq, "
         "sparse gradient bilan ham yaxshi ishlaydi.")

    doc.add_paragraph()
    heading(doc, "6.5. Mini-Batch o'qitish", level=2)
    body(doc,
         "   Batch size = 512\n"
         "   Epochlar soni = 80\n"
         "   Har epochda ma'lumotlar aralashtirilib mini-batch'larga bo'linadi.\n"
         "\n"
         "   20,000 namuna / 512 ≈ 39 ta mini-batch har epochda.\n"
         "   Jami yangilanishlar: 80 × 39 = 3,120 ta Adam qadami.")

    doc.add_paragraph()
    heading(doc, "6.6. L2 Regularizatsiya", level=2)
    body(doc,
         "   λ = 1e-5\n"
         "\n"
         "   L2 regularizatsiya og'irliklarning juda katta qiymatlar olishini "
         "cheklaydi (weight decay). Bu overfitting'ning oldini oladi:\n"
         "   ∂L/∂W += λ × W")

    # ══════════════════════════════════════════════════════════════════════════
    # 7. FRAUD QANDAY ANIQLANADI
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "7. Fraud Qanday Aniqlanadi — Bosqichlar")
    body(doc,
         "Har bir yangi tranzaksiya kelganda quyidagi to'liq pipeline ishlaydi:")

    doc.add_paragraph()
    simple_table(doc,
        ["#", "Bosqich", "Nima qilinadi", "Taxminiy vaqt"],
        [
            ["1", "So'rov qabul qilindi",
             "API yoki Streamlit orqali JSON qabul qilinadi. "
             "Kanal, format va majburiy maydonlar tekshiriladi.", "5–15 ms"],
            ["2", "Karta/Hisob tekshiruvi",
             "Karta tarmog'i, hisob holati va kunlik limit tekshiriladi.", "10–30 ms"],
            ["3", "Qurilma va Lokatsiya tahlili",
             "device va location risky_device, risky_location bayroqlariga aylantiriladi. "
             "Geo-IP ma'lumotlari qo'shiladi.", "15–40 ms"],
            ["4", "Xulq-atvor tahlili",
             "Foydalanuvchi profili bilan solishtiriladi. "
             "Amount normalizatsiya qilinadi. 7 ta interaction feature hisoblanadi.", "20–50 ms"],
            ["5", "DeepFraudNet modeli",
             "25 ta EXT_FEATURES vektori transformer orqali o'tadi. "
             "Forward pass: Input→256→128→64→Sigmoid. "
             "P(fraud) ∈ [0, 1] qaytariladi.", "15–35 ms"],
            ["6", "Threshold baholash",
             "P(fraud) optimal threshold bilan solishtiriladi. "
             "ALLOW / REVIEW / BLOCK qaror aniqlanadi. "
             "Risk signal barlari hisoblanadi.", "5–10 ms"],
            ["Jami", "—", "—", "< 100 ms"],
        ],
        col_widths=[0.4, 1.8, 3.5, 1.2]
    )

    doc.add_paragraph()
    heading(doc, "7.1. Threshold va Qaror Qabul Qilish", level=2)
    body(doc,
         "Optimal threshold o'qitishdan keyin F1-Score maksimallashtirib aniqlanadi "
         "(0.05 dan 0.95 gacha 91 ta qiymat sinab ko'riladi).")
    body(doc,
         "   P(fraud) < threshold           →  ALLOW   (tranzaksiya o'tkaziladi)\n"
         "   threshold ≤ P(fraud) < 1.6×threshold →  REVIEW   (operator tekshiradi)\n"
         "   P(fraud) ≥ max(1.6×threshold, 0.70)   →  BLOCK    (darhol bloklanadi)")

    doc.add_paragraph()
    heading(doc, "7.2. Risk Signal Engine", level=2)
    body(doc,
         "Har tranzaksiya uchun 5 ta risk signali hisoblanib, "
         "vizual progress bar sifatida ko'rsatiladi:")
    simple_table(doc,
        ["Signal", "Past risk (0–30%)", "O'rta risk (30–70%)", "Yuqori risk (70–100%)"],
        [
            ["Qurilma",   "ios, android, chrome, safari",
             "linux",     "emulator, unknown"],
            ["Lokatsiya", "Toshkent, Samarqand, Buxoro va b.",
             "Istanbul, Dubai, Almaty",  "foreign_ip, unknown"],
            ["Summa",     "Normalized < 0.55",
             "Normalized 0.55–0.88",    "Normalized > 0.88"],
            ["Xulq-atvor", "Freq yuqori, oxirgi tx yaqin",
             "Chastota o'rta",           "Birinchi tx, yuqori miqdor"],
            ["Vaqt",      "08:00–22:00 (ish soatlari)",
             "22:00–00:00",              "00:00–05:59 (tun)"],
        ],
        col_widths=[1.4, 1.9, 1.9, 2.2]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 8. METRIKALAR
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "8. Model Natijalari va Metrikalar")
    body(doc,
         "Model 40,000 ta namunada o'qitildi va 10,000 ta test namunasida baholandi. "
         "Quyida haqiqiy o'lchovlar keltirilgan:")

    doc.add_paragraph()
    simple_table(doc,
        ["Metrika", "Qiymat", "Ta'rifi"],
        [
            ["ROC-AUC",   "0.857",
             "Fraud/Normal ajratish qobiliyati. 1.0 = mukammal, 0.5 = tasodifiy"],
            ["F1-Score",  "0.454",
             "Precision va Recall o'rtacha harmonikasi — muvozanatli ko'rsatkich"],
            ["Precision", "0.524",
             "BLOCK/REVIEW diyilganlarning qancha ulushi haqiqiy fraud"],
            ["Recall",    "0.400",
             "Barcha haqiqiy fraudning qancha ulushi aniqlandi"],
            ["Accuracy",  "~95%",
             "Barcha bashoratlarning to'g'riligi (sinf nomutanosibligi tufayli yanglituvchi)"],
            ["Threshold", "~0.66",
             "Optimal F1 erishiladigan chegara (o'qitishda aniqlanadi)"],
        ],
        col_widths=[1.4, 1.1, 4.9]
    )

    doc.add_paragraph()
    body(doc, "Nima uchun Recall muhimroq?", bold=True)
    body(doc,
         "Fraud tizimida False Negative (model 'normal' dedi, aslida fraud) xatolik "
         "False Positive (model 'fraud' dedi, aslida normal) dan ancha qimmat. "
         "Chunki o'tib ketgan fraud = moliyaviy yo'qotish. "
         "Shuning uchun Recall ko'rsatkichi optimallashtirish maqsadi hisoblanadi.")

    doc.add_paragraph()
    body(doc, "ROC-AUC 0.857 nima anglatadi?", bold=True)
    body(doc,
         "Tasodifiy tanlangan fraud tranzaksiyasi tasodifiy tanlangan normal tranzaksiyadan "
         "85.7% hollarda yuqori fraud ehtimolligini oladi. "
         "Sanoat standarti 0.85–0.92 oralig'ida. "
         "Model ushbu diapazonning pastki chegarasida joylashgan — "
         "chunki dataset ataylab 'shovqinli' qilib generatsiya qilingan.")

    doc.add_paragraph()
    body(doc, "Metrikalarni oshirish yo'llari (kelajak):", bold=True)
    bullet(doc, "Ko'proq namuna (50K → 100K) — Recall +3–5%")
    bullet(doc, "LSTM qatlam — vaqtiy ketma-ketlikni o'rganish")
    bullet(doc, "Xgboost/LightGBM ensemble — ROC-AUC +2–4%")
    bullet(doc, "Hyperparameter tuning — learning rate, batch size, qatlamlar soni")

    # ══════════════════════════════════════════════════════════════════════════
    # 9. DEMO PANELI
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "9. Streamlit Demo Paneli")
    body(doc,
         "SafeNet interaktiv demo paneli Streamlit orqali qurilgan va GitHub'da ochiq: "
         "github.com/XafizadinovUsnatdin/Fraudetection")

    doc.add_paragraph()
    body(doc, "Panel bo'limlari:", bold=True)
    simple_table(doc,
        ["Tab", "Nomi", "Funksionalligi"],
        [
            ["1", "Model o'qitish",
             "Ma'lumotlar yuklash, model o'qitish progress bar bilan, natijalar ko'rsatish"],
            ["2", "Tranzaksiyani tekshirish",
             "Bitta tranzaksiyani qo'lda kiritib fraud ehtimolini aniqlash"],
            ["3", "Model ko'rsatkichlari",
             "ROC egri chizig'i, PR egri chizig'i, Confusion Matrix, Feature Importance"],
            ["4", "Ma'lumotlar",
             "Dataset statistikasi, fraud taqsimoti, vizualizatsiyalar"],
            ["5", "Jonli simulyatsiya",
             "100 ta tranzaksiya oqimi real-vaqtda ko'rsatiladi, 4 ta stsenariy"],
        ],
        col_widths=[0.5, 2.2, 4.7]
    )

    doc.add_paragraph()
    body(doc, "Jonli simulyatsiya ssenariylari:", bold=True)
    simple_table(doc,
        ["Stsenariy", "Tavsif", "Fraud ulushi (taxm.)"],
        [
            ["Normal ish kuni",
             "Odatiy kunlik tranzaksiyalar, turli vaqt va summalar",
             "~5%"],
            ["Tun fraud hujumi",
             "00:00–05:00 orasidagi tranzaksiyalar, ko'p emulator va chet el IP",
             "~40%"],
            ["Karta klonlash",
             "Bir foydalanuvchi qisqa vaqt ichida ko'p joydan to'lov",
             "~30%"],
            ["Muvozanatli test",
             "50/50 normal va fraud, barcha turdagi qurilma/lokatsiya",
             "~50%"],
        ],
        col_widths=[2.0, 3.5, 1.9]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 10. TEXNIK STACK
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "10. Texnik Stack")
    simple_table(doc,
        ["Komponent", "Texnologiya", "Versiya / Izoh"],
        [
            ["ML Framework",  "NumPy / Pandas",
             "Sklearn yo'q — Streamlit Cloud bilan mos"],
            ["Web Framework", "Streamlit",
             "Interaktiv dashboard, cache_resource, session_state"],
            ["Vizualizatsiya","Plotly",           "Interaktiv grafiklar, gauge chart"],
            ["Seriallash",    "Pickle",            "Model artifact saqlash/yuklash"],
            ["Prezentatsiya", "python-pptx",       "15 slayd PPTX generatsiya"],
            ["Hujjatlashtirish","python-docx",     "Ushbu Word hujjat generatsiyasi"],
            ["Versiya boshqaruvi","Git / GitHub",  "github.com/XafizadinovUsnatdin/Fraudetection"],
            ["Deploy",        "Streamlit Cloud",   "Streamlit.io orqali bepul hosting"],
            ["Til",           "Python 3.11+",      "Type hints, f-strings, walrus operator"],
        ],
        col_widths=[1.7, 1.8, 3.9]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 11. XULOSA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    heading(doc, "11. Xulosa")
    body(doc,
         "SafeNet loyihasi chuqur o'qitish (deep learning) asosida qurilgan "
         "real-vaqt moliyaviy firibgarlikni aniqlash tizimini namoyish etadi. "
         "Asosiy yutuqlar:")

    doc.add_paragraph()
    numbered(doc, "DeepFraudNet (256→128→64 MLP) pure numpy'da implementatsiya qilindi — "
                  "Streamlit Cloud bilan to'liq mos.")
    numbered(doc, "7 ta interaction feature (qurilma×lokatsiya, VPN×qora_ro'yxat va b.) "
                  "ROC-AUC ni oshirishga hissa qo'shdi.")
    numbered(doc, "sqrt class weighting (4.36x) oddiy n/(2*n_pos) = 10x dan yaxshiroq — "
                  "precision va recall muvozanatlashdi.")
    numbered(doc, "Risk signal engine 5 ta omil bo'yicha vizual tushuntirish beradi.")
    numbered(doc, "Jonli simulyatsiya 4 ta real stsenariy bilan model xulq-atvorini ko'rsatadi.")
    numbered(doc, "ROC-AUC 0.857 — sanoat standarti (0.85–0.92) doirasida.")

    doc.add_paragraph()
    body(doc,
         "Loyiha kelajakda LSTM yoki Transformer asosida yangilanishi, "
         "O'zbekiston bank tizimlariga REST API orqali integratsiya qilinishi mumkin.",
         color_hex="444444")

    divider(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer.add_run(
        "SafeNet Anti-Fraud AI  |  github.com/XafizadinovUsnatdin/Fraudetection  |  2026")
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    build()
